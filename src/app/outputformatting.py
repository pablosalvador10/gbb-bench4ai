import io
import json

import streamlit as st
from docx import Document

from my_utils.ml_logging import get_logger

# Set up logger
logger = get_logger()


def markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """
    Converts markdown text to a docx document.

    :param markdown_text: The input markdown text to be converted.

    :return: The generated docx document as a BytesIO object.
    """
    doc = Document()
    lines = markdown_text.split("\n")

    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("##### "):
            doc.add_heading(line[6:], level=5)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="ListBullet")
            process_bold_text(line[2:], paragraph)
        else:
            paragraph = doc.add_paragraph()
            process_bold_text(line, paragraph)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def process_bold_text(text: str, paragraph) -> None:
    """
    Processes bold text in markdown and adds it to a docx paragraph.

    :param text: The input text that may contain markdown bold syntax.
    :param paragraph: The docx paragraph to add the processed text to.

    :return: None
    """
    while "**" in text:
        start_bold = text.find("**")
        end_bold = text.find("**", start_bold + 2)
        if start_bold != -1 and end_bold != -1:
            paragraph.add_run(text[:start_bold])
            paragraph.add_run(text[start_bold + 2 : end_bold]).bold = True
            text = text[end_bold + 2 :]
        else:
            break
    paragraph.add_run(text)


def download_chat_history() -> None:
    """
    Provide a button to download the chat history as a JSON file.
    """
    chat_history_json = json.dumps(st.session_state.messages, indent=2)
    st.download_button(
        label="📜 Download Chat",
        data=chat_history_json,
        file_name="chat_history.json",
        mime="application/json",
        key="download-chat-history",
    )


def download_ai_response_as_docx_or_pdf() -> None:
    """
    Provide options to download the AI response as a DOCX or PDF file.
    """
    try:
        doc_io = markdown_to_docx(st.session_state.ai_response)
        file_format = st.selectbox("Select file format", ["DOCX", "PDF"])

        if file_format == "DOCX":
            st.download_button(
                label="📁 Download .docx",
                data=doc_io,
                file_name="AI_Generated_Guide.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download-docx",
            )
        elif file_format == "PDF":
            st.download_button(
                label="📁 Download .pdf",
                data=doc_io,
                file_name="AI_Generated_Guide.pdf",
                mime="application/pdf",
                key="download-pdf",
            )
    except Exception as e:
        logger.error(f"Error generating {file_format} file: {e}")
        st.error(
            f"❌ Error generating {file_format} file. Please check the logs for more details."
        )
